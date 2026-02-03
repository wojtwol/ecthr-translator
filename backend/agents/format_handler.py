"""Format Handler - DOCX parsing and reconstruction."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Lazy import to avoid circular dependencies
_citation_detector = None

def get_citation_detector():
    """Lazy initialization of CitationDetector."""
    global _citation_detector
    if _citation_detector is None:
        from agents.citation_detector import CitationDetector
        _citation_detector = CitationDetector()
    return _citation_detector


class FormatHandler:
    """Handles DOCX format extraction and reconstruction with ECtHR formatting guidelines."""

    def _clean_and_format_text(self, text: str) -> str:
        """
        Czyści i formatuje tekst zgodnie z wytycznymi ETPCz.

        Stosuje:
        1. Eliminuje podwójne/potrójne spacje
        2. Usuwa spacje przed interpunkcją
        3. Stosuje twarde spacje przy sierotach (jednoliterowe spójniki, jednostki)
        4. Zamienia niepoprawne cudzysłowy na polskie
        5. Zamienia łączniki na półpauzy w przedziałach liczbowych

        Args:
            text: Tekst do wyczyszczenia

        Returns:
            Wyczyszczony tekst z twardymi spacjami
        """
        import re

        if not text:
            return text

        # 1. Usuń podwójne/potrójne spacje
        text = re.sub(r'\s{2,}', ' ', text)

        # 2. Usuń spacje przed interpunkcją
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)

        # 3. Zamień cudzysłowy na poprawne polskie pary „..."
        # Normalize all quote variants to ASCII "
        for qch in ['\u201c', '\u201d', '\u201e', '\u201f']:
            text = text.replace(qch, '"')
        # Alternate: odd = opening „ (U+201E), even = closing " (U+201D)
        parts = text.split('"')
        if len(parts) > 1:
            result = parts[0]
            for idx, part in enumerate(parts[1:], 1):
                if idx % 2 == 1:
                    result += '\u201e' + part  # „ opening
                else:
                    result += '\u201d' + part  # " closing
            text = result

        # 4. Zamień łączniki na półpauzy w przedziałach liczbowych
        # §§ 34-45 → §§ 34–45
        text = re.sub(r'(§§?\s*\d+)-(\d+)', r'\1–\2', text)
        # art. 5-8 → art. 5–8
        text = re.sub(r'(art\.\s*\d+)-(\d+)', r'\1–\2', text)
        # paragrafy 10-15 → paragrafy 10–15
        text = re.sub(r'(paragraf[ya]?\s+\d+)-(\d+)', r'\1–\2', text)

        # 5. Twarde spacje przy sierotach
        # Twarda spacja w Unicode: \u00A0 (non-breaking space)
        NBSP = '\u00A0'

        # Jednoliterowe spójniki: i, w, z, o, a, u (na końcu słowa lub na początku)
        for conjuction in ['i', 'w', 'z', 'o', 'a', 'u']:
            # Na początku: "w sprawie" → "w\u00A0sprawie"
            text = re.sub(rf'\b{conjuction}\s+', f'{conjuction}{NBSP}', text)

        # Jednostki i skróty: r., §, art., ust., lit., nr, zob.
        # regex_pattern → display_form (to avoid backslash leaking into output)
        units = [
            ('r\\.', 'r.'), ('§', '§'), ('art\\.', 'art.'), ('ust\\.', 'ust.'),
            ('lit\\.', 'lit.'), ('nr', 'nr'), ('zob\\.', 'zob.'), ('par\\.', 'par.'),
        ]
        for regex_unit, display_unit in units:
            # Przed jednostką: "2020 r." → "2020\u00A0r."
            text = re.sub(rf'(\d+)\s+{regex_unit}', lambda m: f'{m.group(1)}{NBSP}{display_unit}', text)
            # Po jednostce: "art. 5" → "art.\u00A05"
            text = re.sub(rf'{regex_unit}\s+', f'{display_unit}{NBSP}', text)

        # Liczby przed jednostkami: "5 marca" → "5\u00A0marca"
        text = re.sub(r'(\d+)\s+(stycznia|lutego|marca|kwietnia|maja|czerwca|lipca|sierpnia|września|października|listopada|grudnia|dni|miesięcy|lat|roku)', rf'\1{NBSP}\2', text)

        return text

    def extract(self, source_path: str) -> Dict[str, Any]:
        """
        Extract text segments from DOCX while preserving format metadata.

        Args:
            source_path: Path to source DOCX file

        Returns:
            Dictionary with segments and document metadata
        """
        try:
            doc = Document(source_path)
            segments = []

            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if not para.text.strip():  # Skip empty paragraphs
                    continue

                segment = {
                    "index": i,
                    "text": para.text,
                    "format": self._extract_paragraph_format(para),
                    "parent_type": "paragraph",
                }
                segments.append(segment)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            if not para.text.strip():
                                continue

                            segment = {
                                "index": len(segments),
                                "text": para.text,
                                "format": self._extract_paragraph_format(para),
                                "parent_type": "table_cell",
                                "table_position": {
                                    "table": table_idx,
                                    "row": row_idx,
                                    "col": col_idx,
                                },
                            }
                            segments.append(segment)

            document_metadata = self._extract_document_metadata(doc)

            logger.info(f"Extracted {len(segments)} segments from {source_path}")

            return {
                "segments": segments,
                "document_metadata": document_metadata,
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise

    def reconstruct(
        self,
        translated_segments: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        output_path: str,
        color_citations: bool = False,
    ) -> str:
        """
        Reconstruct DOCX from translated segments.

        Args:
            translated_segments: List of translated segments with format metadata
            metadata: Document metadata
            output_path: Path to save reconstructed DOCX
            color_citations: If True, color segments containing case citations (default: False)

        Returns:
            Path to reconstructed DOCX file
        """
        try:
            doc = Document()

            # Apply document-level styles (simplified for Sprint 1)
            # Full style restoration will be in Sprint 2

            # Group segments by type for reconstruction
            table_segments = {}

            # Initialize citation detector if needed
            citation_detector = get_citation_detector() if color_citations else None
            citation_count = 0

            # Reconstruct segments
            for segment in translated_segments:
                # Get translated text with fallbacks
                translated_text = segment.get("target_text") or segment.get("translated_text") or segment.get("text", "")

                # Check for citations in SOURCE text (original document)
                has_citations = False
                if citation_detector:
                    source_text = segment.get("text", "")  # Original source text
                    citations = citation_detector.detect_citations(source_text)
                    has_citations = citations.get("total", 0) > 0
                    if has_citations:
                        citation_count += 1
                        logger.debug(f"Segment {segment.get('index', '?')} contains {citations['total']} citation(s)")

                parent_type = segment.get("parent_type", "paragraph")

                if parent_type == "paragraph":
                    # Wyczyść i sformatuj tekst zgodnie z wytycznymi ETPCz
                    cleaned_text = self._clean_and_format_text(translated_text)
                    para = doc.add_paragraph()
                    para.text = cleaned_text
                    self._apply_paragraph_format(para, segment.get("format", {}), has_citations=has_citations)

                elif parent_type == "table_cell":
                    # Store table segments for later reconstruction
                    table_pos = segment.get("table_position", {})
                    table_idx = table_pos.get("table", 0)

                    if table_idx not in table_segments:
                        table_segments[table_idx] = []

                    # Add citation flag to segment
                    segment["_has_citations"] = has_citations
                    table_segments[table_idx].append(segment)

            # Reconstruct tables (simplified - creates new tables)
            for table_idx in sorted(table_segments.keys()):
                segments_in_table = table_segments[table_idx]
                if not segments_in_table:
                    continue

                # Find max row/col to create table
                max_row = max(s.get("table_position", {}).get("row", 0) for s in segments_in_table)
                max_col = max(s.get("table_position", {}).get("col", 0) for s in segments_in_table)

                # Create table
                table = doc.add_table(rows=max_row + 1, cols=max_col + 1)
                table.style = 'Table Grid'

                # Fill table cells
                for segment in segments_in_table:
                    table_pos = segment.get("table_position", {})
                    row = table_pos.get("row", 0)
                    col = table_pos.get("col", 0)
                    translated_text = segment.get("target_text") or segment.get("translated_text") or segment.get("text", "")
                    has_citations = segment.get("_has_citations", False)

                    if row < len(table.rows) and col < len(table.rows[row].cells):
                        cell = table.rows[row].cells[col]
                        # Clear existing paragraphs and add new one with proper formatting
                        # Wyczyść i sformatuj tekst zgodnie z wytycznymi ETPCz
                        cleaned_text = self._clean_and_format_text(translated_text)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        para.text = cleaned_text
                        if has_citations:
                            # Apply citation color to table cell text
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 140, 0)  # Orange color

            # Save document
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)

            if color_citations and citation_count > 0:
                logger.info(f"Reconstructed DOCX saved to {output_path} (colored {citation_count} segments with citations)")
            else:
                logger.info(f"Reconstructed DOCX saved to {output_path}")

            return output_path

        except Exception as e:
            logger.error(f"Error reconstructing DOCX: {e}")
            raise

    def _extract_paragraph_format(self, para) -> Dict[str, Any]:
        """Extract formatting information from a paragraph."""
        format_data = {
            "alignment": str(para.alignment) if para.alignment else None,
            "style": para.style.name if para.style else None,
        }

        # Extract run-level formatting from first run (simplified)
        if para.runs:
            first_run = para.runs[0]
            format_data["font"] = {
                "name": first_run.font.name,
                "size": first_run.font.size.pt if first_run.font.size else None,
                "bold": first_run.font.bold,
                "italic": first_run.font.italic,
                "underline": first_run.font.underline,
            }

        return format_data

    def _apply_paragraph_format(self, para, format_data: Dict[str, Any], has_citations: bool = False) -> None:
        """Apply formatting to a paragraph."""
        # Apply alignment
        if format_data.get("alignment"):
            try:
                para.alignment = eval(format_data["alignment"])
            except:
                pass

        # Apply style
        if format_data.get("style"):
            try:
                para.style = format_data["style"]
            except:
                pass

        # Apply font formatting to runs
        if para.runs and format_data.get("font"):
            font_data = format_data["font"]
            for run in para.runs:
                if font_data.get("name"):
                    run.font.name = font_data["name"]
                if font_data.get("size"):
                    run.font.size = Pt(font_data["size"])
                if font_data.get("bold") is not None:
                    run.font.bold = font_data["bold"]
                if font_data.get("italic") is not None:
                    run.font.italic = font_data["italic"]
                if font_data.get("underline") is not None:
                    run.font.underline = font_data["underline"]

        # Apply citation color if needed (overrides other colors)
        if has_citations and para.runs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 140, 0)  # Orange color for citations

    def _extract_document_metadata(self, doc) -> Dict[str, Any]:
        """Extract document-level metadata."""
        metadata = {
            "core_properties": {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
            },
            "styles": [style.name for style in doc.styles],
        }
        return metadata
